import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    doc = Document()

    # ページ設定 (余白)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # 全体スタイル設定
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Yu Gothic'
    style_normal.font.size = Pt(10.5)
    style_normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

    lines = md_text.split('\n')
    in_code_block = False
    code_lines = []

    for line in lines:
        stripped = line.strip()

        # コードブロック
        if stripped.startswith('```') or stripped.startswith('````'):
            if in_code_block:
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(6)
                code_text = '\n'.join(code_lines)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_lines = []
            else:
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # 空行
        if not stripped:
            continue

        # 水平線
        if stripped in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            continue

        # 見出し1 (#)
        if stripped.startswith('# '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(text)
            run.font.name = 'Yu Gothic'
            run.font.size = Pt(17)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
            continue

        # 見出し2 (##)
        if stripped.startswith('## '):
            text = stripped[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(text)
            run.font.name = 'Yu Gothic'
            run.font.size = Pt(13.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
            continue

        # 見出し3 (###)
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(9)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text)
            run.font.name = 'Yu Gothic'
            run.font.size = Pt(11.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
            continue

        # 引用 (> )
        if stripped.startswith('> '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
            continue

        # 箇条書き (* or -)
        if re.match(r'^[\*\-]\s+', stripped):
            text = re.sub(r'^[\*\-]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2.5)
            add_formatted_text(p, text)
            continue

        # 番号付きリスト (1. 2. etc)
        if re.match(r'^\d+\.\s+', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(2.5)
            add_formatted_text(p, text)
            continue

        # 通常段落
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4.5)
        p.paragraph_format.line_spacing = 1.2
        add_formatted_text(p, stripped)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")

def add_formatted_text(paragraph, text):
    tokens = re.split(r'(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.font.bold = True
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        elif token.startswith('[') and '](' in token and token.endswith(')'):
            m = re.match(r'\[(.*?)\]\((.*?)\)', token)
            if m:
                run = paragraph.add_run(f"{m.group(1)} ({m.group(2)})")
                run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
                run.font.underline = True
            else:
                paragraph.add_run(token)
        else:
            paragraph.add_run(token)

if __name__ == '__main__':
    md_file = r"C:\Users\ryoma\.gemini\antigravity\brain\69d3f3d0-55d2-4d4a-abf0-7fc63da13177\app_introduction_and_manual.md"
    docx_file1 = r"D:\antigravity2\Cloud-Bookreader\Cloud_BookReader_Manual.docx"
    docx_file2 = r"C:\Users\ryoma\.gemini\antigravity\brain\69d3f3d0-55d2-4d4a-abf0-7fc63da13177\Cloud_BookReader_Manual.docx"
    
    convert_md_to_docx(md_file, docx_file1)
    convert_md_to_docx(md_file, docx_file2)
